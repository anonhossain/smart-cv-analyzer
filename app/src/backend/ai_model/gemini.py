import os
import pdfplumber
import google.generativeai as genai
from docx import Document
from dotenv import load_dotenv
from backend.core.config import settings
from prompt import candidate_match_prompt, generate_questions_prompt, hr_match_prompt

# Configuration
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

class BaseAI:
    """Shared AI utilities to avoid redundancy."""
    @staticmethod
    def get_gemini_response(prompt):
        model = genai.GenerativeModel(os.getenv("MODEL"))
        response = model.generate_content(prompt)
        return response.text

    @staticmethod
    def extract_text_with_pdfplumber(file_path):
        """Unified PDF extraction logic."""
        with pdfplumber.open(file_path) as pdf:
            return "".join(page.extract_text() or "" for page in pdf.pages)

    @staticmethod
    def load_text_file(file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

#----------------------------------------------------------------
# REGULAR CANDIDATE ORCHESTRATOR
#----------------------------------------------------------------

class RegularCandidate:
    def __init__(self):
        self.cv_path = settings.CANDIDATE_CV_PATH
        self.jd_path = settings.CANDIDATE_JD_PATH

    def extract_data(self):
        self.resume_text = BaseAI.extract_text_with_pdfplumber(self.cv_path)
        self.job_desc = BaseAI.load_text_file(self.jd_path)

    def get_analysis(self):
        prompt = candidate_match_prompt(self.job_desc, self.resume_text)
        self.result = BaseAI.get_gemini_response(prompt)

    def run_RegularCandidate(self):
        """Orchestration Pattern: Execute all steps."""
        self.extract_data()
        self.get_analysis()
        return self.result

#----------------------------------------------------------------
# HR RANKER ORCHESTRATOR
#----------------------------------------------------------------

class HRPersonRanker:
    def __init__(self):
        self.cv_dir = settings.HR_CV_DIR
        self.jd_path = settings.HR_JD_DIR

    def extract_jd(self):
        self.job_desc = BaseAI.load_text_file(self.jd_path)

    def process_bulk_ranking(self):
        pdf_files = [f for f in os.listdir(self.cv_dir) if f.lower().endswith('.pdf')]
        self.rankings = {}

        for pdf_file in pdf_files:
            path = os.path.join(self.cv_dir, pdf_file)
            resume_text = BaseAI.extract_text_with_pdfplumber(path)
            prompt = hr_match_prompt(self.job_desc, resume_text)
            
            response = BaseAI.get_gemini_response(prompt).strip()
            # Extract digits only for robust scoring
            score = "".join(filter(str.isdigit, response))
            self.rankings[pdf_file] = int(score) if score else 0
        
        self.rankings = dict(sorted(self.rankings.items(), key=lambda x: x[1], reverse=True))

    def run_HRPersonRanker(self):
        """Orchestration Pattern: Execute all steps."""
        self.extract_jd()
        self.process_bulk_ranking()
        return self.rankings

#----------------------------------------------------------------
# HR QUESTION GENERATOR ORCHESTRATOR
#----------------------------------------------------------------

class HRQuestionGenerator:
    def __init__(self):
        self.cv_dir = settings.HR_CV_DIR
        self.jd_path = settings.HR_JD_DIR
        self.output_dir = settings.OUTPUT_DIR

    def prepare_environment(self):
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
        self.job_desc = BaseAI.load_text_file(self.jd_path)

    def generate_and_save(self):
        pdf_files = [f for f in os.listdir(self.cv_dir) if f.lower().endswith('.pdf')]
        
        for filename in pdf_files:
            resume_path = os.path.join(self.cv_dir, filename)
            resume_text = BaseAI.extract_text_with_pdfplumber(resume_path)
            
            prompt = generate_questions_prompt(self.job_desc, resume_text)
            questions = BaseAI.get_gemini_response(prompt)
            
            self._save_as_docx(questions, filename)

    def _save_as_docx(self, text, original_name):
        base_name = os.path.splitext(original_name)[0]
        output_path = os.path.join(self.output_dir, f"{base_name}_questions.docx")
        
        doc = Document()
        doc.add_heading(f"Interview Questions - {base_name}", 0)
        for line in text.split('\n'):
            doc.add_paragraph(line)
        doc.save(output_path)
        print(f"Generated: {output_path}")

    def run_HRQuestionGenerator(self):
        """Orchestration Pattern: Execute all steps."""
        self.prepare_environment()
        self.generate_and_save()