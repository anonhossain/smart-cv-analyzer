import os
import re
import PyPDF2 as pdf
import pdfplumber
import docx
from dotenv import load_dotenv
import fitz
import google.generativeai as genai
import pdfplumber
import csv
from PyPDF2 import PdfReader
from docx import Document
from docx2pdf import convert
import shutil
import csv

from backend.prompt import candidate_match_prompt, generate_questions_prompt, hr_match_prompt

# Set up Google API key and configure Generative AI
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

#----------------------------------------------------------------
#GEMINI FOR CANDIDATE
#----------------------------------------------------------------

class Gemini:
    CANDIDATE_CV_FILE = os.getenv('CANDIDATE_CV_FILE')  # Path to the candidate's resume PDF
    CANDIDATE_JD_FILE = os.getenv('CANDIDATE_JD_FILE')  # Path to the candidate's job description text file
    MODEL = os.getenv("MODEL")

    @staticmethod
    def get_gemini_response(prompt):
        model = genai.GenerativeModel(Gemini.MODEL)
        response = model.generate_content(prompt)
        return response.text

    @staticmethod
    def extract_text_from_pdf(resume_file_path):
        """Extract text from the uploaded PDF resume."""
        with open(resume_file_path, 'rb') as file:
            reader = pdf.PdfReader(file)
            text = ""
            for page in range(len(reader.pages)):
                page_text = reader.pages[page].extract_text()
                text += str(page_text)
        return text

    @staticmethod
    def load_job_description(job_description_file_path):
        """Load job description from the text file."""
        with open(job_description_file_path, 'r') as file:
            return file.read()

    @staticmethod
    def process_resume():
        job_desc = Gemini.load_job_description(Gemini.CANDIDATE_JD_FILE)
        resume_text = Gemini.extract_text_from_pdf(Gemini.CANDIDATE_CV_FILE)
        
        """Process the resume based on the user's selected action."""
        
        prompt = candidate_match_prompt(job_desc, resume_text)
        
        response_text = Gemini.get_gemini_response(prompt)
        return response_text

class GeminiHR:
    HR_CV_FILE = os.getenv('HR_CV_FILE')  # Directory where the HR CVs are stored
    HR_JD_FILE = os.getenv('HR_JD_FILE')  # Path to the HR job description text file

    @staticmethod
    def get_gemini_response(prompt):
        """Get response from Gemini AI."""
        model = genai.GenerativeModel(GeminiHR.MODEL)
        response = model.generate_content(prompt)
        return response.text

    @staticmethod
    def extract_pdf_text(file_path):
        """Extract text from PDF file."""
        with pdfplumber.open(file_path) as pdf:
            return "".join(page.extract_text() for page in pdf.pages)

    @staticmethod
    def extract_docx_text(file_path):
        """Extract text from DOCX file."""
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    @staticmethod
    def process_resume():
        """Process the resumes based on the user's selected action."""
        
        job_desc = Gemini.load_job_description(GeminiHR.HR_JD_FILE)

        pdf_files = [
            file for file in os.listdir(GeminiHR.HR_CV_FILE)
            if file.lower().endswith('.pdf')
        ]

        percentage_mapping = {}

        for pdf_file in pdf_files:
            pdf_path = os.path.join(GeminiHR.HR_CV_FILE, pdf_file)

            resume_text = Gemini.extract_text_from_pdf(pdf_path)
            prompt = hr_match_prompt(job_desc, resume_text)
            response = Gemini.get_gemini_response(prompt).strip()

            try:
                percentage = int(response.replace("\n", "").strip())
            except ValueError:
                percentage = 0 
            
            percentage_mapping[pdf_file] = percentage
        percentage_mapping = dict(sorted(percentage_mapping.items(), key=lambda item: item[1], reverse=True))
        return percentage_mapping


#----------------------------------------------------------------
# Generate questions based on resumes and job description
#----------------------------------------------------------------

class HR_question_generator:
    HR_CV_FILE = os.getenv('HR_CV_FILE')  # Directory where the HR CVs are stored
    HR_JD_FILE = os.getenv('HR_JD_FILE')  # Path to the HR job description text file
    OUTPUT_DIR = os.getenv('OUTPUT_DIR')
    MODEL = os.getenv("MODEL")

    @staticmethod
    def get_gemini_response(prompt):
        model = genai.GenerativeModel(HR_question_generator.MODEL)
        response = model.generate_content(prompt)
        return response.text

    @staticmethod
    def extract_text_from_pdf(path):
        with open(path, 'rb') as file:
            reader = pdf.PdfReader(file)
            return ''.join([page.extract_text() for page in reader.pages if page.extract_text()])

    @staticmethod
    def load_job_description(path):
        with open(path, 'r') as file:
            return file.read()

    @staticmethod
    def save_text_as_pdf(text, output_path):
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=12)
        doc.save(output_path)
        doc.close()

    @staticmethod
    def save_text_as_docx(text, output_path):
        doc = Document()
        for line in text.split('\n'):
            doc.add_paragraph(line)
        doc.save(output_path)

    @staticmethod
    def process_resumes():
        if not os.path.exists(HR_question_generator.OUTPUT_DIR):
            os.makedirs(HR_question_generator.OUTPUT_DIR)

        job_desc = HR_question_generator.load_job_description(HR_question_generator.HR_JD_FILE)

        for filename in os.listdir(HR_question_generator.HR_CV_FILE):
            if filename.lower().endswith('.pdf'):
                resume_path = os.path.join(HR_question_generator.HR_CV_FILE, filename)
                resume_text = HR_question_generator.extract_text_from_pdf(resume_path)
                prompt = generate_questions_prompt(job_desc, resume_text)
                questions = HR_question_generator.get_gemini_response(prompt)
                # Create filenames
                base_name = os.path.splitext(filename)[0]
                docx_output_path = os.path.join(HR_question_generator.OUTPUT_DIR, base_name + "_questions.docx")
                # Save both versions
                HR_question_generator.save_text_as_docx(questions, docx_output_path)
                print(f"Generated: {base_name}questions.docx")